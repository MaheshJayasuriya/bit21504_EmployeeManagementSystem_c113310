#!/usr/bin/env python3
"""
Generate a submission-ready .docx from the repository documentation markdown.

Behavior:
- Prefer `documentation/SUBMISSION.md` as the source (this file was generated earlier).
- Fallback to `documentation/Instructions.md` if SUBMISSION.md is missing.
- Convert simple Markdown headings, paragraphs, bullet lists and fenced code blocks to a Word document.
- Output: `documentation/Submission.docx`

Usage:
1) Create a virtual environment (recommended) and install dependency:
   python -m pip install -r documentation/requirements.txt
2) Run the script from the repository root:
   python documentation/generate_submission.py

The script is intentionally simple and aimed at producing a clean, human-readable Word document for submission.
"""

import os
import sys
import datetime
from docx import Document
from docx.shared import Pt

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC_CANDIDATES = [
    os.path.join(ROOT, 'documentation', 'SUBMISSION.md'),
    os.path.join(ROOT, 'documentation', 'Instructions.md'),
]

def find_source():
    for p in SRC_CANDIDATES:
        if os.path.exists(p):
            return p
    return None


def md_to_docx(src_path, out_path):
    with open(src_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    # Title and metadata
    doc.core_properties.title = 'Assignment Submission'
    doc.add_heading('Assignment Submission', level=1)
    doc.add_paragraph(f'Generated: {datetime.datetime.now().isoformat()}')
    doc.add_paragraph(f'Source file: {os.path.relpath(src_path)}')
    doc.add_paragraph('')

    in_code = False
    code_lines = []

    for raw in lines:
        line = raw.rstrip('\n')
        stripped = line.strip()

        # fenced code block toggling
        if stripped.startswith('```'):
            in_code = not in_code
            if not in_code:
                # flush code block
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                # small spacer after code
                doc.add_paragraph('')
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        # headings
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=3)
            continue

        # bullet lists (simple)
        if stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
            continue

        # numbered lists (simple)
        if stripped and (stripped[0].isdigit() and stripped.split('.', 1)[0].isdigit()):
            # treat as normal paragraph but with numbering marker removed
            after = stripped.split('.', 1)
            if len(after) > 1:
                doc.add_paragraph(after[1].strip(), style='List Number')
                continue

        # blank line -> paragraph break
        if stripped == '':
            doc.add_paragraph('')
            continue

        # inline code backticks: keep as normal text but preserve backticks
        doc.add_paragraph(line)

    # final save
    out_dir = os.path.dirname(out_path)
    os.makedirs(out_dir, exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    src = find_source()
    if not src:
        print('No source Markdown found. Expected SUBMISSION.md or Instructions.md in documentation/.', file=sys.stderr)
        sys.exit(2)

    out = os.path.join(os.path.dirname(src), 'Submission.docx')
    print(f'Reading {src} -> building {out}')
    try:
        path = md_to_docx(src, out)
        print('Wrote', path)
    except Exception as e:
        print('Error while generating docx:', e, file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
